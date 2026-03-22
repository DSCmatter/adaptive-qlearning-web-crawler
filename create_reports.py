import os
from docx import Document

def add_md_to_doc(doc, md_filepath):
    if os.path.exists(md_filepath):
        with open(md_filepath, 'r', encoding='utf-8') as f:
            content = f.read()
            doc.add_paragraph(f"--- Content from {os.path.basename(md_filepath)} ---")
            doc.add_paragraph(content)
    else:
        print(f"File not found: {md_filepath}")

def generate_reports():
    phases = [1, 2, 3, 4, 5]
    reports_dir = 'docs/reports'
    os.makedirs(reports_dir, exist_ok=True)
    
    for phase in phases:
        doc = Document()
        doc.add_heading(f'Phase {phase} Report', 0)
        
        md_file = f'docs/phases/PHASE_{phase}.md'
        add_md_to_doc(doc, md_file)
        
        res_file = f'data/results/PHASE_{phase}_RESULTS.md'
        add_md_to_doc(doc, res_file)
        
        out_path = os.path.join(reports_dir, f'PHASE_{phase}.docx')
        doc.save(out_path)
        print(f"Generated {out_path}")

if __name__ == '__main__':
    generate_reports()